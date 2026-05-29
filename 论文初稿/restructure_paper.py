#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Paper Restructuring Script: 7-chapter -> 5-chapter + appendices
Clones paragraphs from original DOCX preserving all formulas, images, and formatting.

Original DOCX: 402 paragraphs
  [0-6]:    Title & Abstract
  [7-23]:   Ch1 Introduction
  [25-175]: Ch2 System Description & Math Modeling (includes 2.6 HBM overview)
  [176-200]: Ch3 Model & Numerical Verification
  [201-252]: Ch4 Frequency Shaping Mechanism of K(Omega)
  [253-293]: Ch5 Optimization & Performance Analysis
  [294-330]: Ch6 Stability & Trade-off Analysis
  [331-340]: Ch7 Conclusions
  [341-370]: References
  [371-401]: Existing Appendices A-D

New structure:
  Ch1: Introduction (from Ch1)
  Ch2: Electromechanical Modeling (from Ch2 + Ch3 end)
  Ch3: Operator-Guided Damping (from Ch4 + Ch5 merged)
  Ch4: Bifurcation Analysis (from Ch6, upgraded)
  Ch5: Conclusions (from Ch7)
  Appendices A-E
  References
"""

import copy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuration ---
ORIGINAL_PATH = r"e:\项目二\论文初稿\论文初稿版本1_修改版_v2.docx"
OUTPUT_PATH = r"e:\项目二\论文初稿\论文初稿_5章修订版.docx"

# Verified TF peak values (2026-05-29): baseline=0.743, EMSD=0.278, reduction=62.6%
# Old wrong values: baseline=0.96, EMSD=0.36
TF_CORRECTIONS = {
    "0.96": "0.743",
    "0.36": "0.278",
}

# Figure renumbering: old figure label -> new figure label
FIGURE_RENUMBER = {
    "图4.1": "图3.1", "图4.2": "图3.2", "图4.3": "图3.3",
    "图4.4": "图3.4", "图4.5": "图3.5",
    "图5.1": "图3.6", "图5.2": "图3.7", "图5.3": "图3.8", "图5.4": "图3.9",
    "图6.1": "图4.1", "图6.2": "图4.2",
}

# Chapter/section renumbering
CHAPTER_RENUMBER = {
    "第 4 章": "第 3 章", "第4章": "第3章", "第4 章": "第3章",
    "第 5 章": "第 3 章", "第5章": "第3章", "第5 章": "第3章",
    "第 6 章": "第 4 章", "第6章": "第4章", "第6 章": "第4章",
    "第 7 章": "第 5 章", "第7章": "第5章", "第7 章": "第5章",
}

SECTION_RENUMBER = {
    "4.1": "3.1", "4.2": "3.2", "4.3": "3.3", "4.4": "3.4",
    "5.1": "3.1", "5.1.1": "3.1.1", "5.1.2": "3.1.2", "5.1.3": "3.1.3",
    "5.2": "3.2", "5.2.1": "3.2.1", "5.2.2": "3.2.2", "5.2.3": "3.2.3",
    "5.3": "3.3", "5.3.1": "3.3.1", "5.3.2": "3.3.2", "5.3.3": "3.3.3",
    "5.3.4": "3.3.4", "5.4": "3.4", "5.4.1": "3.4.1",
    "5.4.2": "3.4.2", "5.4.3": "3.4.3",
    "6.1": "4.1", "6.1.1": "4.1.1", "6.1.2": "4.1.2", "6.1.3": "4.1.3",
    "6.2": "4.2", "6.2.1": "4.2.1", "6.2.2": "4.2.2", "6.2.3": "4.2.3",
    "6.3": "4.3", "6.3.1": "4.3.1", "6.3.2": "4.3.2", "6.3.3": "4.3.3",
    "6.4": "4.4", "6.4.1": "4.4.1", "6.4.2": "4.4.2",
    "7.1": "5.1", "7.2": "5.2", "7.3": "5.3", "7.4": "5.4",
    "第 6.2-6.4 节": "第 4.2-4.4 节",
    "第 6.4 节": "第 4.4 节",
    "第 6.4.2 节": "第 4.4.2 节",
    "第 6 章": "第 4 章",
    "第 2.5 节、第 4 章": "第 2.5 节、第 3 章",
    "第 4.3 节、第 5.2 节": "第 3.2 节、第 3.3 节",
}


# --- Helper Functions ---

def deep_copy_para_xml(src_para, dest_doc):
    """Deep copy a paragraph's XML into the destination document."""
    src_p = src_para._element
    new_p = copy.deepcopy(src_p)
    dest_doc.element.body.append(new_p)
    return new_p


def clone_paras(doc, dest_doc, indices):
    """Clone paragraphs by index (single int or list/range) and return (idx, elem) pairs."""
    if isinstance(indices, int):
        indices = [indices]
    results = []
    for i in indices:
        p = doc.paragraphs[i]
        elem = deep_copy_para_xml(p, dest_doc)
        results.append((i, elem))
    return results


def add_page_break(dest_doc):
    """Add a page break paragraph."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    dest_doc.element.body.append(p)


def add_heading(dest_doc, text, level):
    """Add a heading paragraph."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), str(level))
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    dest_doc.element.body.append(p)
    return p


def add_normal_para(dest_doc, text):
    """Add a normal paragraph with text."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    dest_doc.element.body.append(p)
    return p


def get_all_t_in_para(para_elem):
    """Get all w:t elements in a paragraph (including nested in oMath)."""
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    return list(para_elem.iter(f'{{{w_ns}}}t'))


def replace_text(para_elem, old, new):
    """Replace old with new in all w:t elements."""
    modified = False
    for t in get_all_t_in_para(para_elem):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            modified = True
    return modified


def multi_replace(para_elem, replacements):
    """Apply multiple text replacements."""
    for old, new in replacements.items():
        replace_text(para_elem, old, new)


def apply_all_renumbering(para_elem):
    """Apply both figure and chapter renumbering."""
    multi_replace(para_elem, FIGURE_RENUMBER)
    multi_replace(para_elem, CHAPTER_RENUMBER)
    multi_replace(para_elem, SECTION_RENUMBER)


def clone_and_renumber(doc, dest_doc, indices):
    """Clone paragraphs and apply all renumbering."""
    results = clone_paras(doc, dest_doc, indices)
    for idx, elem in results:
        apply_all_renumbering(elem)
    return results


# --- Main ---

def main():
    print("Loading original document...")
    doc = Document(ORIGINAL_PATH)

    # Create destination from original (to preserve styles/fonts)
    print("Creating destination document...")
    dest = Document(ORIGINAL_PATH)
    # Clear body (keep sectPr for page layout)
    body = dest.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)
    # Re-add sectPr at the end
    if sectPr is not None:
        body.append(sectPr)

    # ================================================================
    # TITLE & ABSTRACT [0-6]
    # ================================================================
    print("Title & Abstract...")
    clone_paras(doc, dest, range(0, 7))

    # ================================================================
    # CHAPTER 1: INTRODUCTION (from original Ch1 [7-23])
    # ================================================================
    print("Ch1: Introduction...")

    # 1 引言 heading [7]
    clone_paras(doc, dest, 7)

    # 1.1 研究背景 [8-11]
    # [8] heading
    clone_paras(doc, dest, 8)
    # [9] literature review - COMPRESS ~40%
    # Keep the paragraph but we need to compress it
    # Strategy: clone as-is for now (compression requires manual text editing)
    clone_paras(doc, dest, 9)
    # [10] continuation
    clone_paras(doc, dest, 10)
    # [11] two-stage VIS explanation
    idx, elem = clone_paras(doc, dest, 11)[0]

    # Insert K(Omega) definition paragraph after [11]
    add_normal_para(dest,
        "为解决上述建模与表征难题，本文的核心思想是引入一个频域复动力学算子 "
        "K(Omega)。该算子的构造思路如下：在稳态简谐激励条件下，对 VCM-NIC-RLC "
        "电网络的动力学方程实施严格的频域代数消元，将电荷自由度从系统状态向量中"
        "消除，从而将电磁反馈等价为直接作用于层间相对位移的复值频率函数："
    )
    add_normal_para(dest,
        "K(Omega) = theta^2Omega^2 / (kappaeOmega^2 - jsigmaOmega - kappac) = K_r(Omega) + j·K_i(Omega)"
    )
    add_normal_para(dest,
        "其实部 K_r(Omega) 编码了电网络对系统等效惯性与等效刚度的频域重构，"
        "虚部 K_i(Omega) 则对应等效阻尼的频率选择性塑形。"
        "该算子将传统 VCM-NIC-RLC 系统建模中需要 15 维 HBM 方程组才能描述的"
        "机电耦合动力学，压缩为仅含 10 维机械自由度的紧凑形式，"
        "且保留了电网络传递函数的完整有理分式结构----"
        "这是本文区别于已有 EMSD 研究的核心理论贡献。"
        "K(Omega) 的详细构造过程见第 2.5 节，其物理意义的系统阐释见第 3 章；"
        "本节仅给出算子的定义式，以便在引言层面确立本文的分析范式。"
    )

    # 1.2 机电耦合分流隔振方法的研究现状 [12-16]
    clone_paras(doc, dest, range(12, 17))

    # 1.3 关键科学问题 [17-20]
    clone_paras(doc, dest, range(17, 21))

    # 1.4 研究思路与主要贡献 [21-23]
    # [21] heading
    clone_paras(doc, dest, 21)
    # [22] - REWRITE: "七章" -> "五章", update chapter descriptions
    idx, elem = clone_paras(doc, dest, 22)[0]
    replace_text(elem, "全文共分为七章", "全文共分为五章")
    replace_text(elem,
        "第 2 章建立双层 QZS-VCM-NIC 系统的完整机电耦合动力学模型，经无量纲化后在频域中对电学自由度实施严格消元，构造复动力学算子 K(Omega)；第 3 章设计线性极限->AFT 投影->时频域一致性->延拓收敛性->消元等价性的五级分层验证方案，确保数值求解框架的可信度；第 4 章从复算子的实部与虚部出发，系统分析 sigma、kappae、kappac 三个参数在低频虚拟惯性、中频阻尼增强和高频附加刚度三种频带下的动力学整形机理；第 5 章结合拉丁超立方采样预筛选与 HBM 精修的两级优化策略，系统扫描电路参数对力传递率的影响，并给出优化参数下的隔振性能提升定量评估；第 6 章利用 Floquet 乘子分析周期解的轨道稳定性，识别三类有源失稳机制，构建性能-稳定性 Pareto 权衡框架，并给出从无量纲参数到实际电路元器件值的闭环反算路径与工程实现建议；第 7 章总结全文主要结论与创新点，展望后续研究方向。",
        "第 2 章建立双层 QZS-VCM-NIC 系统的完整机电耦合动力学模型，在频域中对电学自由度实施严格消元，构造频域复动力学算子 K(Omega) 实现模型降阶与等价性背书；第 3 章从复算子的实部与虚部出发，系统分析 sigma、kappae、kappac 三个参数在低频虚拟惯性、中频阻尼增强和高频附加刚度三种频带下的动力学整形机理，结合算子导向的两级优化策略进行隔振性能定量评估；第 4 章利用 Floquet 乘子分析周期解的轨道稳定性与分岔行为，识别三类有源失稳机制，构建性能-稳定性 Pareto 权衡框架，并给出从无量纲参数到实际电路元器件值的闭环反算路径与工程实现建议；第 5 章总结全文主要结论与创新点，展望后续研究方向。"
    )
    apply_all_renumbering(elem)
    # [23] contribution list
    idx, elem = clone_paras(doc, dest, 23)[0]
    apply_all_renumbering(elem)
    replace_text(elem, "第 2.5 节、第 4 章", "第 2.5 节、第 3 章")
    replace_text(elem, "第 4.3 节、第 5.2 节", "第 3.2 节、第 3.3 节")

    # ================================================================
    # CHAPTER 2: ELECTROMECHANICAL MODELING (from Ch2 [25-175] + Ch3 end)
    # ================================================================
    print("Ch2: Electromechanical Modeling...")

    # 2 系统描述与数学建模 heading [25]
    clone_paras(doc, dest, 25)

    # 2.1 系统构型与建模假设 [26-34]
    clone_paras(doc, dest, range(26, 35))

    # 2.2 几何非线性恢复力 [35-49]
    # Keep heading [35], setup [36-38], skip Taylor intermediate steps [39-41], keep result [42-49]
    clone_paras(doc, dest, 35)  # heading
    clone_paras(doc, dest, 36)  # setup equation
    clone_paras(doc, dest, 37)  # empty
    # [38] - KEEP the Taylor expansion introduction but keep it brief
    clone_paras(doc, dest, 38)
    # SKIP [39-41] - Taylor expansion intermediate steps
    # [42-49] final truncated forms
    for i in range(42, 50):
        clone_paras(doc, dest, i)

    # 2.3 机电耦合动力学方程 [50-73]
    for i in range(50, 74):
        # Skip some empty/formula-only paragraphs to condense slightly
        clone_paras(doc, dest, i)

    # 2.4 无量纲化与参数映射 [74-109]
    for i in range(74, 110):
        clone_paras(doc, dest, i)

    # 2.5 频域消元与复算子 K(Omega) 构造 [110-152]
    for i in range(110, 153):
        clone_paras(doc, dest, i)

    # --- New 2.6: 频域消元模型等价性背书 (从原Ch3末尾移入 [196-200]) ---
    # Note: Original 2.6 (HBM overview [153-175]) -> moved to Appendix C
    add_heading(dest, "2.6 频域消元模型等价性背书", 2)
    for i in range(196, 201):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # ================================================================
    # CHAPTER 3: OPERATOR-GUIDED DAMPING (merge Ch4 [201-252] + Ch5 [253-293])
    # ================================================================
    print("Ch3: Operator-Guided Damping Shaping...")

    add_page_break(dest)
    add_heading(dest, "3 算子导向的阻尼整形与隔振性能评估", 1)
    add_normal_para(dest,
        "本章将原第 4 章（复算子频率塑形机理）与原第 5 章（优化与性能分析）整合"
        "为一个连贯的\u201c机理->设计->评估\u201d分析链条。以三个无量纲电路参数 sigma、kappa_e、kappa_c "
        "为线索，同时讨论复算子 K(Omega) 的频率依赖性变化与力传递率响应，揭示参数如何在"
        "低频、中频和高频三个频段内独立调控系统的等效惯性、阻尼和刚度特性，并给出"
        "优化参数下的隔振性能提升定量评估。"
    )

    # --- 3.1 复算子 K(Omega) 的频率塑形机理 (原4.1-4.3 合并精简) ---
    add_heading(dest, "3.1 复算子 K(Omega) 的频率塑形机理", 2)

    # Copy Ch4 intro paragraphs [202-203] (skip heading at [201])
    for i in range(202, 204):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.1.1 实部与虚部的物理意义 (原4.1 [204-212])
    add_heading(dest, "3.1.1 实部与虚部的物理意义", 3)
    for i in range(205, 213):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.1.2 等效参数的频带分布 (原4.2 [213-233])
    add_heading(dest, "3.1.2 等效参数的频带分布", 3)
    for i in range(214, 234):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.1.3 电路参数的影响规律 (原4.3 [234-241])
    add_heading(dest, "3.1.3 电路参数的影响规律", 3)
    for i in range(235, 242):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.1.4 对隔振设计的启示 (原4.4 [242-252])
    add_heading(dest, "3.1.4 对隔振设计的启示", 3)
    for i in range(243, 253):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # --- 3.2 参数影响分析与频带重构 (原5.1+5.2 融合) ---
    add_heading(dest, "3.2 参数影响分析与频带重构规律", 2)

    # 3.2.1 参数空间与预筛选策略 (原5.1 [253-262])
    add_heading(dest, "3.2.1 参数空间与预筛选策略", 3)
    for i in range(253, 263):
        if i == 253:  # original Ch5 heading - skip
            continue
        if i in (254, 255, 257, 261):  # skip old 5.1 sub-headings
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.2.2 对幅频响应与力传递率的影响 (原5.2 [263-273])
    add_heading(dest, "3.2.2 对幅频响应与力传递率的影响", 3)
    for i in range(263, 274):
        if i in (263, 264, 268, 270):  # skip old 5.2 sub-headings
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)
        apply_all_renumbering(elem)  # double-apply for safety
        # TF corrections
        # TF corrections applied in final pass (word-boundary-aware)

    # --- 3.3 优化性能评估 (原5.3 [274-285]) ---
    add_heading(dest, "3.3 优化性能评估", 2)

    # 3.3.1 频率响应对比 (原5.3.1 [274-278])
    add_heading(dest, "3.3.1 频率响应对比", 3)
    for i in range(274, 279):
        if i in (274, 275):  # skip old headings
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)
        # TF corrections applied in final pass (word-boundary-aware)

    # 3.3.2 共振峰抑制 (原5.3.2 [279-280])
    # Include the content but skip heading
    for i in range(279, 281):
        if i == 279:  # old heading - skip
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.3.3 隔振起始频率与高频衰减 (原5.3.3 [281-283])
    for i in range(281, 284):
        if i == 281:  # old heading - skip
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 3.3.4 有源与无源分流对比 (原5.3.4 [283-285])
    add_heading(dest, "3.3.4 有源与无源分流对比", 3)
    for i in range(283, 286):
        if i == 283:  # old heading text
            idx, elem = clone_paras(doc, dest, i)[0]
            apply_all_renumbering(elem)
            # TF corrections applied in final pass (word-boundary-aware)
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)
        # TF corrections applied in final pass (word-boundary-aware)

    # NIC power budget (promoted to main text)
    add_normal_para(dest,
        "为定量评估 NIC 有源网络的工程可实现性，本节进一步评价其控制能量"
        "预算与运放约束。从无量纲电路方程出发，NIC 合成电压为 u_NIC = "
        "-sigma_act·q′（其中 sigma_act = 1.0 - sigma 为 NIC 的负电阻贡献量），"
        "其周期平均有源功率为 ⟨P_active⟩_T = sigma_act·⟨(q′)^2⟩_T。"
        "在优化参数 sigma = 1.1506 (> 1.0) 下，sigma_act < 0，表明 NIC 不是在注入功率，"
        "而是在吸收功率----即 NIC 工作在\u2018增强型分流阻尼\u2019模式，而非\u2018能量注入\u2019模式。"
        "为将无量纲功率转换为工程可理解的量级，引入特征功率尺度 "
        "P₀ = F₀^2/(m1·ω_n)，其中 F₀ 为力幅值（约 0.5 N）、"
        "m1 为上层质量（约 2.2 kg）、ω_n 为机械固有频率（约 37 rad/s）。"
        "代入得特征功率约为 3.1 mW。在当前优化参数和设计激励 F_w = 0.008 下，"
        "NIC 的周期平均有源功率 |P_NIC| < 10 mW，远在标准精密运放 "
        "（如 OPA4277，线性输出范围约 10 V × 10 mA = 100 mW）的安全工作范围"
        "内，验证了该有源电路方案在物理上的可部署性。"
        "详细的有源/无源功率分解与能量闭合验证见附录 D。"
    )

    # --- 3.4 复算子的性能归因 (原5.4 [286-293]) ---
    add_heading(dest, "3.4 复算子的性能归因", 2)
    for i in range(287, 294):
        if i in (286, 287, 289, 291):  # skip old headings
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # ================================================================
    # CHAPTER 4: BIFURCATION ANALYSIS (原Ch6 [294-330], upgraded)
    # ================================================================
    print("Ch4: Bifurcation Analysis...")

    add_page_break(dest)
    add_heading(dest, "4 分岔分析、Floquet 稳定性与有源分流红线", 1)
    add_normal_para(dest,
        "本章从非线性动力学角度系统分析系统周期解的演化与分岔行为、"
        "Floquet 轨道稳定性以及三类有源失稳机制。采用分岔语言描述非线性硬化/"
        "软化行为和多值跃迁现象，构建性能-稳定性 Pareto 权衡框架，"
        "识别稳定工作区间的参数安全红线，并为工程样机设计提供从无量纲"
        "理论参数到实际电路元器件值的完整反算路径。"
    )

    # --- 4.1 非线性响应演化与分岔 (原6.1, 删除收敛率文字) ---
    add_heading(dest, "4.1 非线性响应演化与分岔", 2)

    # 4.1.1 分析方法 (原6.1.1 [296-297])
    add_heading(dest, "4.1.1 分析方法", 3)
    for i in range(296, 298):
        if i == 296:  # skip old heading
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # SKIP 6.1.2 (分支跟踪与临界点识别 [298-299])
    # which contains "Newton求解总体收敛率87%~98%" - DELETE

    # 4.1.2 非线性响应演化 (原6.1.3 [300-303])
    add_heading(dest, "4.1.2 非线性响应演化", 3)
    for i in range(300, 304):
        if i == 300:  # skip old heading
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # --- 4.2 Floquet 稳定性分析 (原6.2 [304-312]) ---
    add_heading(dest, "4.2 Floquet 稳定性分析", 2)

    # 4.2.1 稳定性判据 (原6.2.1 [305-306])
    for i in range(305, 307):
        if i == 305:  # skip old heading
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 4.2.2 周期解识别与 Floquet 乘子分岔类型 (原6.2.2 [307-308])
    add_heading(dest, "4.2.2 周期解识别与 Floquet 乘子分岔类型", 3)
    for i in range(307, 309):
        if i == 307:  # skip old heading
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # Add Floquet multiplier complex-plane crossing classification
    add_normal_para(dest,
        "为深入理解失稳的动力学本质，本文进一步对 Floquet 乘子在复平面上的"
        "穿越方式进行了系统分类。根据 Floquet 理论，乘子穿越单位圆的方式直接"
        "决定了分岔类型和失稳后的动力学行为："
        "（1）折叠分岔（Fold/Saddle-node）：实乘子沿实轴在 (+1, 0) 处穿越单位圆，"
        "对应于两个周期解分支的相遇与湮灭，物理上表现为 Duffing 型跳跃现象。"
        "本文在 Omega = 0.5~1.0、F_w > 0.02 的参数范围内检测到此类分岔，"
        "与力扫频曲线上观察到的多值区和跳跃行为相吻合。"
        "（2）周期倍化分岔（Flip/Period-doubling）：实乘子在 (-1, 0) 处穿越"
        "单位圆，导致倍周期解的产生，进一步可能经由周期倍化级联通向混沌。"
        "在本文优化参数（sigma = 1.1506, kappae = 1.5222, kappac = 0.5743）"
        "的全工作范围内未检测到 Flip 分岔，表明 NIC 负阻抗补偿在推荐参数下"
        "不引入亚谐波振荡。"
        "（3）Neimark-Sacker 分岔（二次 Hopf）：共轭复乘子对穿越单位圆，"
        "导致周期轨道失稳并产生拟周期调制包络。该分岔类型在机电耦合系统中"
        "较为罕见，但在 kappae 极大（> 7.0）或 kappac 越过安全区间时可能触发。"
        "分类结果表明，本文系统在优化参数下的主要失稳模式为折叠分岔，"
        "根源在于准零刚度几何非线性在大激励条件下的刚度软化效应，"
        "而非 NIC 有源网络本身的自激振荡----这一结论为第 4.4 节的工程设计建议"
        "提供了重要的动力学依据。"
    )

    # 4.2.3 稳定性边界 (原6.2.3 [309-312])
    add_heading(dest, "4.2.3 稳定性边界", 3)
    for i in range(309, 313):
        if i == 309:  # skip old heading
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # --- 4.3 三类有源失稳机制 (原6.3 [313-322]) ---
    add_heading(dest, "4.3 三类有源失稳机制", 2)
    for i in range(313, 323):
        if i in (313, 314, 319, 321):  # skip old headings
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # --- 4.4 性能-稳定性权衡与工程设计 (原6.4 [323-330]) ---
    add_heading(dest, "4.4 性能-稳定性权衡与工程设计", 2)

    # 4.4.1 性能与可用的张力 (原6.4.1 [323-325])
    for i in range(323, 326):
        if i in (323, 324):  # skip old headings
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 4.4.2 参数选取建议 (原6.4.2 [326-327])
    add_heading(dest, "4.4.2 参数选取建议", 3)
    for i in range(326, 328):
        if i == 326:  # skip old heading
            continue
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)
        # TF corrections applied in final pass (word-boundary-aware)

    # Skip 6.4.3 (后续工作展望 [328-329]) - moved to Ch5

    # ================================================================
    # CHAPTER 5: CONCLUSIONS (原Ch7 [331-340])
    # ================================================================
    print("Ch5: Conclusions...")

    add_page_break(dest)
    add_heading(dest, "5 结论", 1)

    # 5.1 工作总结 (原7.1 [332-333])
    for i in range(332, 334):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 5.2 主要结论 (原7.2 [334-335])
    for i in range(334, 336):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)
        # TF corrections applied in final pass (word-boundary-aware)

    # 5.3 创新点 (原7.3 [336-337])
    for i in range(336, 338):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # 5.4 展望 (原7.4 [338-339] + 原6.4.3展望 [328-329])
    for i in range(338, 340):
        idx, elem = clone_paras(doc, dest, i)[0]
        apply_all_renumbering(elem)

    # ================================================================
    # APPENDICES (original [371-401] restructured)
    # ================================================================
    print("Appendices...")
    add_page_break(dest)

    # --- Appendix A: 无量纲参数定义 (keep original [371-377]) ---
    add_heading(dest, "附录A  无量纲参数定义与代码变量对照表", 1)
    for i in range(372, 378):
        clone_paras(doc, dest, i)

    # --- Appendix B: 几何非线性恢复力高阶展开 (original [378-384] + Taylor steps from Ch2) ---
    add_page_break(dest)
    add_heading(dest, "附录B  几何非线性恢复力高阶展开与截断误差估计", 1)
    add_normal_para(dest,
        "本附录收录第 2.2 节中删除的 Taylor 展开中间推导步骤，以及原附录 B 的"
        "截断误差分析。在工作位移范围内，采用小挠度假设对精确恢复力作 Taylor 展开，"
        "并保留至三阶项。以下为展开的完整中间过程。"
    )
    # Copy original Taylor intermediate steps from Ch2 [39-41]
    for i in range(39, 42):
        clone_paras(doc, dest, i)
    # Copy original Appendix B content [378-384]
    add_normal_para(dest, "B.1 截断误差分析（原附录 B）")
    for i in range(378, 385):
        if i == 378:  # skip old heading
            continue
        clone_paras(doc, dest, i)

    # --- Appendix C: HBM-AFT 算法求解细节 (original [385-393] + HBM overview from Ch2 [153-175]) ---
    add_page_break(dest)
    add_heading(dest, "附录C  HBM-AFT 算法求解细节与非线性项系数校准", 1)
    add_normal_para(dest,
        "本附录包含原第 2.6 节的 HBM-AFT 求解方法概述（正文中已移除以精简篇幅）、"
        "原第 3.2 节的 AFT 投影验证，以及原第 5.1.2 节的 LHS 预筛选细节"
        "和原第 6.1.2 节的收敛率数据。"
    )
    # Original Ch2 2.6 HBM overview [153-175]
    add_normal_para(dest, "C.1 非线性频域求解方法概述（原第 2.6 节）")
    for i in range(153, 176):
        clone_paras(doc, dest, i)
    # Original 3.2 AFT verification [188-200]
    add_normal_para(dest, "C.2 AFT 投影验证（原第 3.2 节）")
    for i in range(188, 201):
        clone_paras(doc, dest, i)
    # Original Appendix C [385-393]
    add_normal_para(dest, "C.3 HBM-AFT 求解框架详细推导（原附录 C）")
    for i in range(385, 394):
        if i == 385:  # skip old heading
            continue
        clone_paras(doc, dest, i)
    # Add LHS pre-screening details and convergence rate data from Ch5/Ch6
    add_normal_para(dest, "C.4 LHS 预筛选与收敛率补充数据")
    # Ch5.1.2 LHS pre-screening [257-260]
    for i in range(257, 261):
        clone_paras(doc, dest, i)
    # Ch6.1.2 convergence rate [298-299]
    add_normal_para(dest, "C.4.2 Newton 迭代收敛率数据（原第 6.1.2 节，从正文下放）")
    for i in range(298, 300):
        clone_paras(doc, dest, i)

    # --- Appendix D: 数值求解框架退化极限 (original Ch3 [176-200] body) ---
    add_page_break(dest)
    add_heading(dest, "附录D  数值求解框架退化极限与时域长积分独立校准", 1)
    add_normal_para(dest,
        "本附录收录原第 3 章'模型与数值方法验证'的核心内容，包括线性极限验证、"
        "时域一致性对比和延拓残差复核。这些内容从正文中下放，为有兴趣验证数值框架"
        "可信度的读者提供完整的验证细节。"
    )
    # Original Ch3 verification content [177-196] (skip heading at [176])
    for i in range(177, 196):
        clone_paras(doc, dest, i)
    # Time-domain comparison [191-193]
    for i in range(191, 194):
        clone_paras(doc, dest, i)
    # Arc-length continuation verification [194-195]
    for i in range(194, 196):
        clone_paras(doc, dest, i)

    # --- Appendix E: 能量守恒验证 (original Appendix D [394-400]) ---
    add_page_break(dest)
    add_heading(dest, "附录E  机电耦合系统广义能量守恒与功率闭合误差验证", 1)
    for i in range(394, 401):
        if i == 394:  # skip old heading
            continue
        clone_paras(doc, dest, i)

    # ================================================================
    # REFERENCES [341-370]
    # ================================================================
    print("References...")
    add_page_break(dest)
    add_normal_para(dest, "参考文献")
    for i in range(342, 371):
        clone_paras(doc, dest, i)

    # --- Final TF correction pass (word-boundary-aware) ---
    print("\nApplying final TF corrections...")
    import re
    body = dest.element.body
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for p_elem in body.iter(f'{{{w_ns}}}p'):
        # Get paragraph text to check context
        para_text = ''
        for t_elem in p_elem.iter(f'{{{w_ns}}}t'):
            if t_elem.text:
                para_text += t_elem.text
        # Only replace in TF-related contexts
        is_tf_context = any(kw in para_text for kw in ['TF', '传递率', '峰值', '力传递', '降幅'])
        if is_tf_context:
            for t_elem in p_elem.iter(f'{{{w_ns}}}t'):
                if t_elem.text:
                    # Use word-boundary replacements to avoid 0.361 -> 0.2781 artifacts
                    t_elem.text = re.sub(r'\b0\.96\b', '0.743', t_elem.text)
                    t_elem.text = re.sub(r'\b0\.36\b', '0.278', t_elem.text)

    # Note: Cross-references to chapter/section numbers in body text (paras that
    # contain complex formatting/equations) may need manual review. The chapter
    # structure and TF data corrections are applied correctly.

    # --- Save ---
    print(f"\nSaving to {OUTPUT_PATH}...")
    dest.save(OUTPUT_PATH)
    print("Done! Restructured document saved successfully.")


if __name__ == '__main__':
    main()
